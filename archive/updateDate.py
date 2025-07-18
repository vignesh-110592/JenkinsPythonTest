from docx import Document
from datetime import datetime
import os

def replace_date_in_run(run, old_date, new_date):
    if old_date in run.text:
        run.text = run.text.replace(old_date, new_date)

def replace_date_preserving_format(doc, old_date, new_date):
    # Check all paragraphs
    for para in doc.paragraphs:
        for run in para.runs:
            replace_date_in_run(run, old_date, new_date)
    
    # Check inside tables as well
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        replace_date_in_run(run, old_date, new_date)

def update_dates_in_files(source_dir, dest_dir, filenames, old_date, new_date):
    for filename in filenames:
        input_path = os.path.join(source_dir, filename)
        output_path = os.path.join(dest_dir, filename)

        if not os.path.exists(input_path):
            print(f"❌ Source file not found: {input_path}")
            continue

        try:
            doc = Document(input_path)
            replace_date_preserving_format(doc, old_date, new_date)
            doc.save(output_path)
            print(f"✅ Updated and saved: {output_path}")
        except Exception as e:
            print(f"❌ Failed to process {input_path}: {e}")

# ----- Jenkins paths -----
source_workspace = "/source"  # Source workspace in Jenkins
destination_workspace = "/dest"  # Destination workspace in Jenkins
files_to_update = ["AnschreibenRaw.docx", "LebenslaufRaw.docx"]

# Placeholder date to be replaced and current date as replacement
old_date = "__DATE__"  # Example of the previously inserted date (can be a placeholder like "##DATE##")
new_date = datetime.now().strftime("%d.%m.%Y")

update_dates_in_files(source_workspace, destination_workspace, files_to_update, old_date, new_date)
