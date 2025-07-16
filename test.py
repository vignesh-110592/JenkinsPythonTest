from docx import Document

def replace_text_in_paragraph(paragraph, old_text, new_text):
    if old_text in paragraph.text:
        inline = paragraph.runs
        for i in range(len(inline)):
            if old_text in inline[i].text:
                inline[i].text = inline[i].text.replace(old_text, new_text)
                

def replace_text_in_doc(file_path, old_text, new_text, output_path=None):
    doc = Document(file_path)

    for para in doc.paragraphs:
        replace_text_in_paragraph(para, old_text, new_text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_text_in_paragraph(para, old_text, new_text)

    output_path = output_path or file_path
    doc.save(output_path)
    print(f"Text replaced and saved to: {output_path}")

if __name__ == "__main__":
    file_path = 'S:/DevOps/JenkinsPipeline/AnschreibenRaw.docx'

    new_company_name = input("Enter the new company name: ").strip()
    output_path = 'S:/DevOps/JenkinsPipeline/Anschreiben.docx'
    output_path = output_path if output_path else None

    replace_text_in_doc(file_path, 'Company_name', new_company_name, output_path)
