import sys
import os
from docx import Document
from datetime import datetime

def verify_placeholders(doc, expected_values):
    text = "\n".join([para.text for para in doc.paragraphs])
    for key, val in expected_values.items():
        if val not in text:
            print(f"Verification failed: '{val}' not found.")
            return False
    return True

if __name__ == "__main__":
    if len(sys.argv) != 4:
        print("Usage: python verify_docx_text.py <CompanyName> <PersonName> <PositionName>")
        sys.exit(1)

    company_name = sys.argv[1]
    person_name = sys.argv[2]
    position_name = sys.argv[3]

    current_date = datetime.today().strftime("%d.%m.%Y")

    expected_anschreiben = {
        "__DATE__": current_date,
        "Company_name": company_name,
        "Person_name": person_name,
        "Position_name": position_name
    }

    expected_lebenslauf = {
        "__DATE__": current_date
    }

    base_path = os.getcwd()

    doc1 = Document(os.path.join(base_path, "Anschreiben.docx"))
    doc2 = Document(os.path.join(base_path, "Lebenslauf.docx"))

    if not verify_placeholders(doc1, expected_anschreiben):
        sys.exit(1)

    if not verify_placeholders(doc2, expected_lebenslauf):
        sys.exit(1)

    print("Verification passed.")
