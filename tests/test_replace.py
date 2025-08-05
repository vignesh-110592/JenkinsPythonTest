import os
import sys
import tempfile
import subprocess
import shutil
from pathlib import Path
from docx import Document
from datetime import datetime

def create_mock_docx(path, paragraphs):
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    doc.save(path)

def test_replace_script_with_args():
    # Create a temp repo root with app/, input/, output/, tests/
    repo_root = Path(tempfile.mkdtemp())
    app_dir = repo_root / "app"
    input_dir = repo_root / "input"
    output_dir = repo_root / "output"
    tests_dir = repo_root / "tests"

    app_dir.mkdir()
    input_dir.mkdir()
    output_dir.mkdir()
    tests_dir.mkdir()

    # Prepare test DOCX input files
    anschreiben_path = input_dir / "AnschreibenRaw.docx"
    lebenslauf_path = input_dir / "LebenslaufRaw.docx"

    create_mock_docx(anschreiben_path, [
        "__DATE__",
        "Company_name",
        "Greeting Person_name, applying for Position_name."
    ])
    create_mock_docx(lebenslauf_path, [
        "__DATE__"
    ])

    # Copy and patch the original script into app/
    repo_root_str = str(repo_root).replace("\\", "\\\\")  # Escape backslashes for Windows paths
    original_script = Path("app/replace_docx_text.py").read_text()
    patched_script = original_script.replace(
        "dirname(dirname(__file__))",
        f'r"{repo_root_str}"'
    )
    script_path = app_dir / "replace_docx_text.py"
    script_path.write_text(patched_script)

    # Arguments
    args = ["OpenAI", "Dear", "Alice", "Engineer"]

    # Run the script via subprocess
    result = subprocess.run(
        [sys.executable, str(script_path)] + args,
        cwd=str(app_dir),
        capture_output=True,
        text=True
    )

    print("STDOUT:\n", result.stdout)
    print("STDERR:\n", result.stderr)
    assert result.returncode == 0, "Script did not exit cleanly"

    # Check Anschreiben output
    anschreiben_output = output_dir / "Anschreiben.docx"
    assert anschreiben_output.exists(), "Anschreiben output not found"

    doc = Document(anschreiben_output)
    doc_text = "\n".join(p.text for p in doc.paragraphs)

    assert "__DATE__" not in doc_text
    assert "Company_name" not in doc_text
    assert "Greeting" not in doc_text
    assert "Person_name" not in doc_text
    assert "Position_name" not in doc_text

    assert "OpenAI" in doc_text
    assert "Dear" in doc_text
    assert "Alice" in doc_text
    assert "Engineer" in doc_text

    # Check Lebenslauf output
    lebenslauf_output = output_dir / "Lebenslauf.docx"
    assert lebenslauf_output.exists(), "Lebenslauf output not found"

    doc2 = Document(lebenslauf_output)
    doc2_text = "\n".join(p.text for p in doc2.paragraphs)

    current_date = datetime.today().strftime("%d.%m.%Y")
    assert current_date in doc2_text
    assert "__DATE__" not in doc2_text

    # Clean up temp repo
    shutil.rmtree(repo_root)
